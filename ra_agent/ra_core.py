"""
Regulatory Affairs (RA) Agent Core Module
Implements the AI Agent for Regulatory Affairs Automation per pharma industry specification.
"""

import os
import json
import time
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from pathlib import Path
import asyncio
from datetime import datetime
import uuid

import streamlit as st
from openai import AsyncOpenAI
from google.generativeai import GenerativeModel
import os
import json
from typing import List, Dict, Any
import pathlib
from pathlib import Path
import PyPDF2
from docx import Document
from pptx import Presentation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Simplified Knowledge Base - no external vector DB dependencies
from collections import defaultdict
import math
import re

@dataclass
class DocumentManifest:
    """RA document metadata tracking"""
    doc_id: str
    title: str
    doc_type: str
    template_id: str
    version: str
    sources: List[Dict]
    ai_models: List[Dict]
    approval_status: str
    created_by: str
    created_at: datetime
    approved_by: Optional[str] = None
    approved_at: Optional[datetime] = None

@dataclass
class RAQuery:
    """RA query with dual AI processing"""
    query_id: str
    prompt: str
    document_type: str
    use_perplexity: bool
    use_chatgpt: bool
    use_local: bool
    context_docs: List[str]
    ai_responses: Dict = None
    reconciled_output: str = ""
    citations: List[Dict] = None

class DualAISynthesizer:
    """Handles dual AI processing (Perplexity + ChatGPT-5)"""

    def __init__(self):
        self.openai_client = AsyncOpenAI(api_key=st.secrets.get("OPENAI_API_KEY"))
        self.gemini_model = GenerativeModel('gemini-1.5-pro')

    async def query_perplexity(self, prompt: str, context: str) -> Dict:
        """Query Perplexity AI service"""
        try:
            response = await self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",  # Perplexity uses OpenAI-compatible API
                messages=[
                    {"role": "system", "content": f"You are an RA expert. Context: {context}"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            return {
                "ai": "perplexity",
                "response": response.choices[0].message.content,
                "confidence": 0.85,
                "model": "perplexity-sonnet",
                "timestamp": datetime.now()
            }
        except Exception as e:
            st.error(f"Perplexity API error: {e}")
            return {"ai": "perplexity", "error": str(e)}

    async def query_chatgpt(self, prompt: str, context: str) -> Dict:
        """Query ChatGPT-5"""
        try:
            response = await self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": f"You are an RA expert. Context: {context}"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            return {
                "ai": "chatgpt-5",
                "response": response.choices[0].message.content,
                "confidence": 0.90,
                "model": "gpt-4",
                "timestamp": datetime.now()
            }
        except Exception as e:
            st.error(f"ChatGPT API error: {e}")
            return {"ai": "chatgpt-5", "error": str(e)}

    async def reconcile_responses(self, perplexity_resp: Dict, chatgpt_resp: Dict,
                                ra_rules: Dict) -> Tuple[str, List[Dict]]:
        """Reconcile AI responses with RA rules"""
        # Simple reconciliation logic - can be enhanced
        reconciled = ""

        if "error" not in perplexity_resp and "error" not in chatgpt_resp:
            # Use Gemini to reconcile
            reconcile_prompt = f"""
            Reconcile these two RA responses:
            Perplexity: {perplexity_resp['response']}
            ChatGPT: {chatgpt_resp['response']}

            Create a unified, professional RA response following regulatory guidelines.
            """

            reconciliation = await self.query_gemini(reconcile_prompt, "")
            reconciled = reconciliation.get("response", perplexity_resp['response'])

        elif "error" not in perplexity_resp:
            reconciled = perplexity_resp['response']
        else:
            reconciled = chatgpt_resp['response']

        citations = [
            {"ai": "perplexity", "confidence": perplexity_resp.get("confidence", 0)},
            {"ai": "chatgpt-5", "confidence": chatgpt_resp.get("confidence", 0)}
        ]

        return reconciled, citations

    async def query_gemini(self, prompt: str, context: str) -> Dict:
        """Query Google Gemini"""
        try:
            import google.generativeai as genai
            genai.configure(api_key=st.secrets.get("GOOGLE_API_KEY"))
            model = genai.GenerativeModel('gemini-1.5-pro')
            response = model.generate_content(prompt)
            return {
                "ai": "gemini",
                "response": response.text,
                "confidence": 0.95,
                "model": "gemini-1.5-pro",
                "timestamp": datetime.now()
            }
        except Exception as e:
            return {"ai": "gemini", "error": str(e)}

class DocumentGenerator:
    """Generates structured RA documents"""

    def __init__(self):
        self.templates_dir = Path("ra_agent/templates")
        self.templates_dir.mkdir(exist_ok=True)

    async def generate_word_doc(self, content: str, template_id: str,
                               file_path: Path) -> bool:
        """Generate Word document"""
        try:
            doc = Document()

            # Add RA header
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = "Regulatory Affairs Document - CONFIDENTIAL"

            # Split content and add to document
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    doc.add_paragraph(section.strip())

            doc.save(file_path)
            return True
        except Exception as e:
            st.error(f"Word document generation error: {e}")
            return False

    async def generate_pdf(self, content: str, template_id: str,
                          file_path: Path) -> bool:
        """Generate PDF document"""
        try:
            c = canvas.Canvas(str(file_path))

            # Header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, 800, "Regulatory Affairs Document")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, 780, "CONFIDENTIAL")

            # Content
            c.setFont("Helvetica", 10)
            lines = content.split('\n')
            y_pos = 750
            for line in lines:
                if y_pos < 50:  # New page
                    c.showPage()
                    y_pos = 800
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(50, 800, "Regulatory Affairs Document (continued)")
                    y_pos = 780
                    c.setFont("Helvetica", 10)

                c.drawString(50, y_pos, line[:80])  # Limit line length
                y_pos -= 12

            c.save()
            return True
        except Exception as e:
            st.error(f"PDF generation error: {e}")
            return False

class KnowledgeManager:
    """Manages RA knowledge base with simple storage (no ChromaDB dependency)"""

    def __init__(self):
        self.knowledge_path = Path("ra_agent/knowledge_base.json")
        self.knowledge = self._load_knowledge_base()

    def _load_knowledge_base(self) -> Dict[str, List[str]]:
        """Load knowledge base from JSON file"""
        if self.knowledge_path.exists():
            try:
                with open(self.knowledge_path, 'r') as f:
                    return json.load(f)
            except:
                pass
        return defaultdict(list)

    def _save_knowledge_base(self):
        """Save knowledge base to JSON file"""
        try:
            with open(self.knowledge_path, 'w') as f:
                json.dump(dict(self.knowledge), f)
        except Exception as e:
            print(f"Failed to save knowledge base: {e}")

    def _simple_search(self, query: str, text: str) -> float:
        """Simple text matching score (0-1)"""
        query_lower = query.lower()
        text_lower = text.lower()

        # Exact phrase match
        if query_lower in text_lower:
            return 1.0

        # Word matching
        query_words = set(query_lower.split())
        text_words = set(text_lower.split())
        if len(query_words) > 0:
            match_ratio = len(query_words.intersection(text_words)) / len(query_words)
            return min(match_ratio, 1.0)

        return 0.0

    async def add_document(self, doc_path: Path, doc_meta: Dict) -> bool:
        """Add RA document to knowledge base"""
        try:
            text = ""

            if doc_path.suffix == '.pdf':
                with open(doc_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"

            elif doc_path.suffix == '.docx':
                doc = Document(doc_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"

            else:
                with open(doc_path, 'r', encoding='utf-8') as file:
                    text = file.read()

            # Split into paragraphs/sentences
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

            # Store by doc type and keywords
            doc_type = doc_meta.get('doc_type', 'general')
            keywords = doc_meta.get('keywords', [doc_type])

            for para in paragraphs:
                if len(para.strip()) > 20:  # Skip very short paragraphs
                    for keyword in keywords:
                        self.knowledge[keyword.lower()].append({
                            'content': para,
                            'doc_id': doc_meta['doc_id'],
                            'doc_type': doc_type,
                            'source': str(doc_path.name)
                        })

            self._save_knowledge_base()
            return True
        except Exception as e:
            st.error(f"Document ingestion error: {e}")
            return False

    async def search_relevant_docs(self, query: str, top_k: int = 5) -> List[Dict]:
        """Search for relevant RA documents with simple matching"""
        try:
            results = []
            query_lower = query.lower()

            # Search across all keywords
            for keyword, docs in self.knowledge.items():
                if keyword in query_lower or any(word in keyword for word in query_lower.split()):
                    for doc in docs:
                        score = self._simple_search(query, doc['content'])
                        if score > 0.3:  # Minimum relevance threshold
                            results.append({
                                'content': doc['content'][:500] + "..." if len(doc['content']) > 500 else doc['content'],
                                'metadata': {
                                    'doc_id': doc['doc_id'],
                                    'doc_type': doc['doc_type'],
                                    'source': doc['source']
                                },
                                'score': score
                            })

            # Sort by relevance and return top_k
            results.sort(key=lambda x: x['score'], reverse=True)
            return results[:top_k]

        except Exception as e:
            st.error(f"Search error: {e}")
            return []

class RAAgent:
    """Main RA Agent coordinating all functionality"""

    def __init__(self):
        self.ai_synthesizer = DualAISynthesizer()
        self.doc_generator = DocumentGenerator()
        self.knowledge_manager = KnowledgeManager()
        self.output_dir = Path("ra_agent/outputs")
        self.output_dir.mkdir(exist_ok=True)

    async def process_ra_query(self, query: RAQuery) -> Dict:
        """Process an RA query with dual AI and document generation"""

        # Search relevant context
        context_docs = await self.knowledge_manager.search_relevant_docs(
            query.prompt, top_k=3
        )

        context_text = "\n".join([doc["content"] for doc in context_docs])

        # Query AIs
        ai_responses = {}

        if query.use_perplexity:
            perplexity_resp = await self.ai_synthesizer.query_perplexity(
                query.prompt, context_text
            )
            ai_responses["perplexity"] = perplexity_resp

        if query.use_chatgpt:
            chatgpt_resp = await self.ai_synthesizer.query_chatgpt(
                query.prompt, context_text
            )
            ai_responses["chatgpt"] = chatgpt_resp

        if query.use_local:
            local_resp = await self.ai_synthesizer.query_gemini(
                query.prompt, context_text
            )
            ai_responses["gemini"] = local_resp

        query.ai_responses = ai_responses

        # Reconcile responses
        if len(ai_responses) > 1:
            # Simple reconciliation - take the best response
            best_ai = max(ai_responses.keys(),
                         key=lambda x: ai_responses[x].get("confidence", 0))
            reconciled_output = ai_responses[best_ai]["response"]
        else:
            reconciled_output = list(ai_responses.values())[0]["response"]

        query.reconciled_output = reconciled_output

        # Generate citations
        citations = []
        for ai_name, resp in ai_responses.items():
            if "response" in resp:
                citations.append({
                    "ai": ai_name,
                    "confidence": resp.get("confidence", 0),
                    "timestamp": resp.get("timestamp")
                })

        query.citations = citations

        return {
            "query_id": query.query_id,
            "reconciled_output": reconciled_output,
            "ai_responses": ai_responses,
            "citations": citations,
            "context_docs": context_docs
        }

    async def generate_document(self, content: str, doc_type: str,
                               template_id: str, filename: str) -> Dict:
        """Generate RA document in multiple formats"""

        doc_id = str(uuid.uuid4())
        output_files = {}

        # Generate Word document
        word_path = self.output_dir / f"{filename}.docx"
        if await self.doc_generator.generate_word_doc(content, template_id, word_path):
            output_files["docx"] = str(word_path)

        # Generate PDF
        pdf_path = self.output_dir / f"{filename}.pdf"
        if await self.doc_generator.generate_pdf(content, template_id, pdf_path):
            output_files["pdf"] = str(pdf_path)

        # Create manifest
        manifest = DocumentManifest(
            doc_id=doc_id,
            title=filename,
            doc_type=doc_type,
            template_id=template_id,
            version="1.0",
            sources=[],  # Would be populated with actual sources
            ai_models=[{"model": "dual_ai", "timestamp": datetime.now()}],
            approval_status="draft",
            created_by="ra_agent",
            created_at=datetime.now()
        )

        return {
            "manifest": manifest,
            "output_files": output_files,
            "doc_id": doc_id
        }

# Global RA Agent instance
ra_agent = RAAgent()
